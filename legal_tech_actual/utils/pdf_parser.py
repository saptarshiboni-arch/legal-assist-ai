import pdfplumber
import docx
import io
import base64
import re
import os
from groq import Groq

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
if not GROQ_API_KEY:
    raise ValueError("GROQ_API_KEY environment variable is not set")
vision_client = Groq(api_key=GROQ_API_KEY)

def extract_text_from_image(file_content: bytes, filename: str) -> str:
    """Extract text from image using Groq Vision model."""
    mime_type = "image/png" if filename.lower().endswith('.png') else "image/jpeg"
    base64_image = base64.b64encode(file_content).decode('utf-8')

    chat_completion = vision_client.chat.completions.create(
        messages=[{
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": (
                        "You are an OCR expert. Carefully read and transcribe ALL the text visible in this contract or document image. "
                        "Include every word, number, clause, and signature block exactly as it appears. "
                        "Return ONLY the raw transcribed text with no extra commentary."
                    )
                },
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:{mime_type};base64,{base64_image}"},
                },
            ],
        }],
        model="meta-llama/llama-4-scout-17b-16e-instruct",
        temperature=0.0,
        max_tokens=4096,
    )
    return chat_completion.choices[0].message.content.strip()


def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """Extract text from PDF, DOCX, JPG, JPEG, or PNG."""
    text = ""
    ext = filename.lower()

    try:
        if ext.endswith('.pdf'):
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        elif ext.endswith('.docx'):
            doc = docx.Document(io.BytesIO(file_content))
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " | "
                    text += "\n"
        elif ext.endswith(('.jpg', '.jpeg', '.png')):
            text = extract_text_from_image(file_content, filename)
        else:
            raise ValueError("Unsupported file format. Please upload PDF, DOCX, JPG, JPEG, or PNG.")

    except Exception as e:
        raise Exception(f"Text extraction failed: {str(e)}")

    text = text.strip()
    if not text or len(text) < 20:
        raise ValueError(
            "No text could be extracted. If this is a scanned PDF, please upload clear JPG/PNG photos of the pages instead."
        )

    return text
